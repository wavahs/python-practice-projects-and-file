from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)
    
document = Document()

# profile picture
document.add_heading("Profile Picture")
document.add_picture("me.jpg", width = Inches(2.0))

# name and contact information
name = input("Enter your full name here: ")
speak(f"Hello {name}, how are you today and what's your phone number?")
phone_number = int(input("Enter phone number here: "))
email = input("Enter your email here: ")        
document.add_paragraph(f"""
                       Name: {name.upper()} \n 
                       Contact: {phone_number} \n
                       Email: {email}""")
                       
# about me
document.add_heading('About me')
document.add_paragraph(input("Tell me about yourself: "))
                       
# work experience
document.add_heading("Work Experience")

company = input("Enter company name: ")
position = input("Enter position held while at that company: ")
from_date = input("From date: ")
to_date = input("To date: ")

while True:
    try:
        from_date <= to_date
        if from_date <= to_date:
            break;
        else:
            print("'To date' must be greater than 'From date', try again")
            to_date = input("To date: ")
            
    except ValueError:
        print("Please enter an integer")
        continue
        
experience_details = input(f"Describe your experience at {company} ")

document.add_paragraph(f"""
Company name: {company.title()} \n
Position held: {position.capitalize()} \n
Period spent: {from_date}-{to_date} \n
Experience details at company: {experience_details.capitalize()} \n""")

while True:
    has_more_experiences = input("Do you have more experience, yes or no? ")
    if has_more_experiences.lower() == "yes":
        company = input("Enter company name: ")
        position = input("Enter position held while at that company: ")
        from_date = input("From date:")
        to_date = input("To date:")
        experience_details = input(f"Describe your experience at {company} ")

        document.add_paragraph(f"""
Company name: {company.title()} \n
Position held: {position.capitalize()} \n
Period spent: {from_date}-{to_date} \n
Experience details at company: {experience_details.capitalize()} \n""")
    else:
        break

document.add_heading("Skills")     
prag = document.add_paragraph(input("Enter skill possessed "))
prag.style = "List Bullet"
    
while True:
    has_more_skills = input("Do you have more skills, yes or no? ")
    if has_more_skills == "yes":
        prag = document.add_paragraph("Enter skill possessed ")
        prag.style = "List Bullet"
    else:
        break
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
document.save("cv1.docx")
